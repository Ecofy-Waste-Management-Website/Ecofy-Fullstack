import docx

def read_template(filepath):
    doc = docx.Document(filepath)
    print("Paragraphs:")
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            print(f"P {i}: {p.text}")
    print("\nTables:")
    for i, table in enumerate(doc.tables):
        print(f"Table {i}:")
        for r_idx, row in enumerate(table.rows):
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.replace('\n', ' ').strip())
            print(f"Row {r_idx}: {row_data}")

if __name__ == '__main__':
    read_template(r"C:\Users\USER\Downloads\Test_Case_Template.docx")
