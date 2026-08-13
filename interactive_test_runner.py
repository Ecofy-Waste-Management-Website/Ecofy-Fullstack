import sys
import importlib
from pathlib import Path
from docx import Document
import traceback

CURRENT_DIR = Path(__file__).resolve().parent
TESTS_DIR = CURRENT_DIR / "Tests"

if str(CURRENT_DIR) not in sys.path:
    sys.path.insert(0, str(CURRENT_DIR))
if str(TESTS_DIR) not in sys.path:
    sys.path.insert(0, str(TESTS_DIR))

import selenium_support

MODULES = [
    "user_auth_module_selenium",
    "service_booking_and_tracking_module_selenium",
    "staff_management_and_admin_module_selenium",
    "payment_module_selenium",
    "chatbot_and_notification_module_selenium",
]

test_results = []
current_test = {}

original_manual_step = selenium_support.manual_step

def hooked_manual_step(message: str) -> None:
    if ":" in message:
        test_id, desc = message.split(":", 1)
        current_test["id"] = test_id.strip()
        current_test["desc"] = desc.strip()
    else:
        current_test["id"] = "Unknown"
        current_test["desc"] = message
        
    print(f"\n--- [ACTION REQUIRED] ---")
    original_manual_step(message)
    print("-------------------------\n")

selenium_support.manual_step = hooked_manual_step

def generate_report(results, output_path):
    doc = Document()
    doc.add_heading("Ecofy Test Execution Report", 0)
    
    for i, res in enumerate(results, 1):
        doc.add_heading(f"{i}) {res.get('id', 'N/A')} - {res.get('name', 'Unknown Test')}", level=1)
        
        p = doc.add_paragraph()
        p.add_run("Test Case ID: ").bold = True
        p.add_run(f"{res.get('id', 'N/A')}\n")
        
        p.add_run("Test Case Title: ").bold = True
        p.add_run(f"{res.get('name', 'Unknown')}\n")
        
        p.add_run("Test Scenario: ").bold = True
        p.add_run(f"{res.get('desc', 'No description captured.')}\n")
        
        p.add_run("Priority: ").bold = True
        p.add_run("High\n")
        
        p.add_run("Test Type: ").bold = True
        p.add_run("Functional/UI\n")
        
        p.add_run("Expected Result: ").bold = True
        p.add_run("The step completes successfully and necessary elements are present.\n")
        
        p.add_run("Actual Result: ").bold = True
        p.add_run(f"{res.get('actual_result', 'N/A')}\n")
        
        p.add_run("Status: ").bold = True
        p.add_run(f"{res.get('status', 'Unknown')}\n")
        
        doc.add_page_break()
        
    doc.save(output_path)
    print(f"\n✅ Report successfully saved to: {output_path}")

def run_all_tests():
    global current_test
    print("\nStarting Interactive Test Runner...")
    print("This will execute all modules sequentially. Please ensure your Debugging Chrome Session is running!")
    
    for module_name in MODULES:
        print(f"\n======================================")
        print(f"LOADING MODULE: {module_name}")
        print(f"======================================")
        try:
            module = importlib.import_module(module_name)
        except Exception as e:
            print(f"Failed to import module {module_name}: {e}")
            continue
            
        driver = selenium_support.build_driver()
        
        # Get all run_* functions in the module
        funcs = [f for f in dir(module) if f.startswith('run_') and callable(getattr(module, f))]
        funcs.sort()
        
        for func_name in funcs:
            func = getattr(module, func_name)
            current_test = {
                "name": func_name,
                "id": "Unknown",
                "desc": "No description captured.",
                "status": "Fail",
                "actual_result": "Failed before completion"
            }
            
            print(f"\n>> Executing Test: {func_name}")
            try:
                func(driver)
                current_test["status"] = "Pass"
                current_test["actual_result"] = "Test completed successfully without errors."
            except Exception as e:
                current_test["status"] = "Fail"
                current_test["actual_result"] = str(e)
                traceback.print_exc()
                print(">> Test FAILED. Continuing to next test...")
            
            # We copy current_test so that references aren't overwritten in the next iteration
            test_results.append(current_test.copy())
            
        selenium_support.shutdown_driver(driver)
        
    report_path = str(Path.home() / "Downloads" / "Test_Case_Report.docx")
    generate_report(test_results, report_path)

if __name__ == "__main__":
    run_all_tests()
